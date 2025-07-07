import sys
import os
import re
import shutil
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def preprocess_image_links(content):
    """Convert patterns like (attachment:xxx:图x-xx.png) to (图x-xx.png)"""
    return re.sub(r'\(attachment:[^:]*:(图[^)]*)\)', r'(\1)', content)

def convert_markdown_to_docx(md_file, template_docx, output_docx):
    """Convert Markdown to Word document using the specified template"""
    # Read the markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Preprocess image links
    content = preprocess_image_links(content)
    
    # Copy template to output file
    shutil.copy(template_docx, output_docx)
    
    # Create a new document using the template
    doc = Document(output_docx)
    
    # Remove all paragraphs except the first one
    for i in range(len(doc.paragraphs)-1, 0, -1):
        p = doc.paragraphs[i]._element
        p.getparent().remove(p)
    
    # Clear the first paragraph if it exists
    if len(doc.paragraphs) > 0:
        doc.paragraphs[0].clear()
    
    # Process the markdown content
    lines = content.split('\n')
    i = 0
    
    # Track state
    in_code_block = False
    code_lines = []
    in_aside_block = False
    aside_lines = []
    
    while i < len(lines):
        line = lines[i]
        
        # Process code blocks
        if line.strip() == '```':
            if not in_code_block:
                # Start of code block
                in_code_block = True
                code_lines = []
            else:
                # End of code block
                in_code_block = False
                # Add the code block content
                for code_line in code_lines:
                    p = doc.add_paragraph(code_line)
                    try:
                        p.style = '代码无行号'
                    except:
                        print(f"Warning: Style '代码无行号' not found in template")
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Process aside blocks
        if line.strip() == '<aside>':
            in_aside_block = True
            aside_lines = []
            i += 1
            continue
        elif line.strip() == '</aside>' and in_aside_block:
            in_aside_block = False
            
            # Process aside content based on first line
            if aside_lines:
                first_line = aside_lines[0]
                style_name = None
                
                if '【避坑指南】' in first_line:
                    style_name = '强提示'
                elif '提示' in first_line:
                    style_name = '提示'
                
                if style_name:
                    for aside_line in aside_lines:
                        # Remove backticks from inline code
                        clean_line = re.sub(r'`([^`]+)`', r'\1', aside_line)
                        p = doc.add_paragraph(clean_line)
                        try:
                            p.style = style_name
                        except:
                            print(f"Warning: Style '{style_name}' not found in template")
            
            i += 1
            continue
        
        if in_aside_block:
            aside_lines.append(line)
            i += 1
            continue
        
        # Process headings
        if line.startswith('# '):
            p = doc.add_paragraph(line[2:])
            try:
                p.style = 'Heading 1'
            except:
                print(f"Warning: Style 'Heading 1' not found in template")
        elif line.startswith('## '):
            p = doc.add_paragraph(line[3:])
            try:
                p.style = 'Heading 2'
            except:
                print(f"Warning: Style 'Heading 2' not found in template")
        elif line.startswith('### '):
            p = doc.add_paragraph(line[4:])
            try:
                p.style = 'Heading 3'
            except:
                print(f"Warning: Style 'Heading 3' not found in template")
        
        # Process numbered lists (using Heading 4 style)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph(line)
            try:
                p.style = 'Heading 4'
            except:
                print(f"Warning: Style 'Heading 4' not found in template")
        
        # Process unordered lists
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:])
            try:
                p.style = '第1级无序列表'
            except:
                print(f"Warning: Style '第1级无序列表' not found in template")
        
        # Process code listings
        elif line.startswith('代码清单'):
            p = doc.add_paragraph(line)
            try:
                p.style = '超强提示标签'
            except:
                print(f"Warning: Style '超强提示标签' not found in template")
        
        # Process table titles
        elif line.startswith('表'):
            p = doc.add_paragraph(line)
            try:
                p.style = '表题'
            except:
                print(f"Warning: Style '表题' not found in template")
        
        # Process table content
        elif line.startswith('|'):
            # This is a table row, we need to handle tables specially
            # Find the table boundaries
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_rows.append(lines[i])
                i += 1
            
            # Parse the table structure
            if len(table_rows) >= 3:  # Header, separator, and at least one data row
                # Create a table with the appropriate number of rows and columns
                cols = len(table_rows[0].split('|')) - 2  # -2 for the empty parts at start/end
                rows = len(table_rows) - 1  # -1 for the separator row
                
                table = doc.add_table(rows=rows, cols=cols)
                table.style = 'Table Grid'
                
                # Fill the table (skip the separator row)
                row_idx = 0
                for table_row in [tr for tr in table_rows if not tr.strip().startswith('|-')]:
                    if row_idx >= rows:
                        break
                    
                    cells = [cell.strip() for cell in table_row.split('|')[1:-1]]
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < cols:
                            table.cell(row_idx, col_idx).text = cell_text
                    
                    row_idx += 1
                
            continue  # Skip incrementing i since we've already moved past the table
        
        # Process warning guidelines
        elif line.startswith('【避坑指南】'):
            p = doc.add_paragraph(line)
            try:
                p.style = '强提示标签'
            except:
                print(f"Warning: Style '强提示标签' not found in template")
        
        # Process tip lines
        elif line.startswith('提示'):
            p = doc.add_paragraph(line)
            try:
                p.style = '提示标签'
            except:
                print(f"Warning: Style '提示标签' not found in template")
        
        # Process images
        elif line.strip().startswith('!['):
            # Extract image information
            match = re.match(r'!\[(.*?)\]\((.*?)\)', line.strip())
            if match:
                img_caption, img_path = match.groups()
                
                # Add image to document
                try:
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_picture(img_path, width=Inches(6.0))  # Default width
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # Add caption
                    caption_p = doc.add_paragraph(img_caption)
                    try:
                        caption_p.style = '图题'
                    except:
                        print(f"Warning: Style '图题' not found in template")
                    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error adding image {img_path}: {e}")
        
        # Process normal text (with inline code handling)
        elif line.strip():
            # Handle inline code
            line_with_inline_code = re.sub(r'`([^`]+)`', r'\1', line)
            
            p = doc.add_paragraph(line_with_inline_code)
            try:
                p.style = 'Normal'
            except:
                print(f"Warning: Style 'Normal' not found in template")
            
            # Apply inline code formatting if needed
            if '`' in line:
                for match in re.finditer(r'`([^`]+)`', line):
                    code_text = match.group(1)
                    for run in p.runs:
                        if code_text in run.text:
                            # Split the run at the inline code
                            parts = run.text.split(code_text)
                            if len(parts) == 2:
                                # Replace the run's text with the first part
                                run.text = parts[0]
                                
                                # Add a new run for the code
                                code_run = p.add_run(code_text)
                                try:
                                    code_run.style = '行内代码'
                                except:
                                    # If style not found, make it bold as fallback
                                    code_run.bold = True
                                
                                # Add a new run for the remainder
                                p.add_run(parts[1])
        
        # Handle empty lines
        else:
            doc.add_paragraph()
        
        i += 1
    
    # Save the document
    doc.save(output_docx)
    print(f"Successfully converted {md_file} to {output_docx}")

def main():
    if len(sys.argv) < 2:
        print("Usage: python converter-by-claude-sonnet-3-7-with-new-prompt.py <markdown_file>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    template_docx = "ch04-to-template.docx"
    output_docx = "template-to-converted-with-new-prompt.docx"
    
    if not os.path.exists(md_file):
        print(f"Error: Input file {md_file} not found")
        sys.exit(1)
    
    if not os.path.exists(template_docx):
        print(f"Error: Template file {template_docx} not found")
        sys.exit(1)
    
    convert_markdown_to_docx(md_file, template_docx, output_docx)

if __name__ == "__main__":
    main()