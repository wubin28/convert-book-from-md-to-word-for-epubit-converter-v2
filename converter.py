#!/usr/bin/env python3
"""
converter: Convert Markdown to Word document with specific styling
"""

import os
import sys
import re
import shutil
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Try importing PIL for image handling
try:
    from PIL import Image
    PIL_AVAILABLE = True
except ImportError:
    PIL_AVAILABLE = False
    print("Warning: PIL/Pillow not available. Image dimensions won't be calculated.")

def main():
    """Main function to convert markdown to docx"""
    if len(sys.argv) != 2:
        print("Usage: python3 converter.py <markdown_file>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    if not os.path.exists(md_file):
        print(f"Error: File {md_file} does not exist")
        sys.exit(1)
    
    # Get the directory of the markdown file
    md_dir = os.path.dirname(md_file) or '.'
    
    # Define output docx file name
    base_name = os.path.splitext(os.path.basename(md_file))[0]
    
    # Special case for template-from.md
    if base_name == "template-from":
        output_file = os.path.join(md_dir, "template-to-converted.docx")
    else:
        output_file = os.path.join(md_dir, f"{base_name}-to-converted.docx")
    
    # Get the template docx file path from the project root directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_file = os.path.join(script_dir, 'template.docx')
    if not os.path.exists(template_file):
        print("Error: Template file 'template.docx' not found in the project root directory")
        sys.exit(1)
    
    # Create a new document from the template (instead of copying it)
    doc = Document(template_file)
    
    # Clear ALL existing content in the document (including tables)
    # First clear tables
    for table in doc.tables[:]:
        tbl = table._element
        tbl.getparent().remove(tbl)
    
    # Then clear paragraphs
    for paragraph in doc.paragraphs[:]:
        p = paragraph._element
        p.getparent().remove(p)
    
    # Read markdown content
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Process markdown content and add to document
    process_markdown(md_content, doc, md_dir)
    
    # Save the document
    doc.save(output_file)
    print(f"Conversion complete! Output file: {output_file}")

def process_markdown(md_content, doc, md_dir):
    """Process markdown content and convert to docx with appropriate styles"""
    # Preprocess image paths: convert (attachment:xxx:图x-xx.png) to (图x-xx.png)
    md_content = re.sub(r'\(attachment:[^:]+:图', r'(图', md_content)
    
    lines = md_content.splitlines()
    i = 0
    
    # Skip the first line if it's a comment
    if lines[0].startswith('<!--'):
        i += 1
        while i < len(lines) and not lines[i].endswith('-->'):
            i += 1
        if i < len(lines):
            i += 1  # Skip the closing comment line
    
    in_code_block = False
    code_content = []
    in_table = False
    table_rows = []
    table_header = False
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End of code block
                in_code_block = False
                
                # Join all code lines
                code_text = '\n'.join(code_content)
                
                # Add code block with appropriate style
                p = doc.add_paragraph()
                p.style = '代码无行号'
                p.add_run(code_text)
                
                code_content = []
            else:
                # Start of code block
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_content.append(line)
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Heading 1
            p = doc.add_paragraph(line[2:], style='Heading 1')
            i += 1
            continue
        
        if line.startswith('## '):
            # Heading 2
            p = doc.add_paragraph(line[3:], style='Heading 2')
            i += 1
            continue
        
        if line.startswith('### '):
            # Heading 3
            p = doc.add_paragraph(line[4:], style='Heading 3')
            i += 1
            continue
            
        # Handle numbered headings (like "1. XXXX")
        numbered_heading_match = re.match(r'^#### (\d+)\. (.+)', line)
        if numbered_heading_match:
            number = numbered_heading_match.group(1)
            content = numbered_heading_match.group(2)
            p = doc.add_paragraph(f"{number}. {content}", style='Heading 4')
            i += 1
            continue
        
        # Handle code listings
        if line.startswith('代码清单'):
            p = doc.add_paragraph(line, style='超强提示标签')
            i += 1
            continue
        
        # Handle unordered lists
        if line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='第1级无序列表')
            i += 1
            continue
        
        # Handle ordered lists with parentheses like （1）
        if re.match(r'^\（\d+\）', line):
            # This is a special type of ordered list with Chinese parentheses
            p = doc.add_paragraph(line, style='Normal')
            i += 1
            continue
        
        # Handle tables
        if line.startswith('表'):
            p = doc.add_paragraph(line, style='表题')
            i += 1
            continue
        
        # Handle table content
        if line.startswith('|'):
            if not in_table:
                in_table = True
                table_rows = [line]
            else:
                table_rows.append(line)
            i += 1
            continue
        elif in_table:
            # End of table
            in_table = False
            
            # Process table
            if len(table_rows) >= 3:  # Need at least header, separator, and one data row
                # First, determine column count from the header
                columns = len(table_rows[0].strip('|').split('|'))
                
                # Create table in document (rows count is table_rows - 1 header row - 1 separator row)
                table = doc.add_table(rows=len(table_rows)-1, cols=columns)
                table.style = 'Table Grid'
                
                # Set table cell style
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            paragraph.style = '表格单元格'
                
                # Process header row first
                header_cells = table_rows[0].strip('|').split('|')
                for col_idx, cell_text in enumerate(header_cells):
                    if col_idx < columns:
                        cell = table.cell(0, col_idx)
                        cell.text = cell_text.strip()
                        # Apply style to each paragraph in the cell
                        for paragraph in cell.paragraphs:
                            paragraph.style = '表格单元格'
                
                # Process data rows (skip header and separator rows)
                for row_idx in range(2, len(table_rows)):
                    cells = table_rows[row_idx].strip('|').split('|')
                    for col_idx, cell_text in enumerate(cells):
                        if col_idx < columns and (row_idx-1) < len(table.rows):
                            cell = table.cell(row_idx-1, col_idx)
                            cell.text = cell_text.strip()
                            # Apply style to each paragraph in the cell
                            for paragraph in cell.paragraphs:
                                paragraph.style = '表格单元格'
            
            # Don't add empty paragraph after table
        
        # Handle tips (避坑指南 and 提示)
        if line.startswith('【避坑指南】'):
            p = doc.add_paragraph(line, style='强提示标签')
            i += 1
            continue
        
        if line.startswith('【提示】'):
            p = doc.add_paragraph(line, style='提示标签')
            i += 1
            continue
        
        # Handle images
        image_match = re.search(r'!\[(.*?)\]\((.*?)\)', line)
        if image_match:
            image_title = image_match.group(1)
            image_path = image_match.group(2)
            
            # Add the image first
            full_image_path = os.path.join(md_dir, image_path)
            if os.path.exists(full_image_path):
                try:
                    # Calculate width for image - use fixed width if PIL not available
                    width_inches = 6.0  # Default width
                    
                    try:
                        if 'PIL_AVAILABLE' in globals() and PIL_AVAILABLE:
                            with Image.open(full_image_path) as img:
                                width, height = img.size
                                # Calculate appropriate width for document (max 6 inches)
                                width_inches = min(width / 96, 6.0)  # Assuming 96 DPI
                    except Exception as e:
                        print(f"Warning: Error calculating image dimensions: {str(e)}")
                    
                    # Add picture
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(full_image_path, width=Inches(width_inches))
                    
                    # Add paragraph for image title below the image
                    p = doc.add_paragraph(image_title, style='图题')
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Warning: Error adding image {image_path}: {str(e)}")
            else:
                print(f"Warning: Image file not found: {image_path}")
                # Still add the title even if image not found
                p = doc.add_paragraph(image_title, style='图题')
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            i += 1
            continue
        
        # Handle aside blocks
        if line.startswith('<aside>'):
            aside_content = []
            i += 1
            while i < len(lines) and not lines[i].startswith('</aside>'):
                aside_content.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # Skip the closing aside tag
                
            # Process aside content based on the type of tip
            for aside_line in aside_content:
                # Skip empty lines and emoji markers
                if aside_line.strip() and not aside_line.strip() == '💡':
                    # Check for inline code in the aside line
                    parts = re.split(r'(`{1,2}.*?`{1,2})', aside_line.strip())
                    
                    if aside_line.startswith('【避坑指南】'):
                        # Use 强提示标签 for the header line
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part.strip('`'))
                                run.style = '行内代码'
                            else:
                                if part.strip():
                                    run = p.add_run(part)
                        p.style = '强提示标签'
                    elif '【避坑指南】' in aside_line:
                        # Use 强提示 for content within 避坑指南 aside
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part.strip('`'))
                                run.style = '行内代码'
                            else:
                                if part.strip():
                                    run = p.add_run(part)
                        p.style = '强提示'
                    elif aside_line.startswith('【提示】'):
                        # Use 提示标签 for the header line
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part.strip('`'))
                                run.style = '行内代码'
                            else:
                                if part.strip():
                                    run = p.add_run(part)
                        p.style = '提示标签'
                    elif '【提示】' in aside_line:
                        # Use 提示 for content within 提示 aside
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part.strip('`'))
                                run.style = '行内代码'
                            else:
                                if part.strip():
                                    run = p.add_run(part)
                        p.style = '提示'
                    else:
                        # For lines within an aside that don't have a marker, use the style of the last marker seen
                        style = '强提示' if any('【避坑指南】' in line for line in aside_content) else '提示'
                        p = doc.add_paragraph()
                        for part in parts:
                            if part.startswith('`') and part.endswith('`'):
                                run = p.add_run(part.strip('`'))
                                run.style = '行内代码'
                            else:
                                if part.strip():
                                    run = p.add_run(part)
                        p.style = style
            continue
            
        # Handle normal paragraphs
        if line.strip():
            # Check for inline code
            parts = re.split(r'(`{1,2}.*?`{1,2})', line)
            if len(parts) > 1:  # Contains inline code
                p = doc.add_paragraph()
                for part in parts:
                    if part.startswith('`') and part.endswith('`'):
                        # Remove the backticks and apply inline code style
                        code_text = part.strip('`')
                        run = p.add_run(code_text)
                        run.style = '行内代码'
                    else:
                        # Regular text
                        if part.strip():
                            run = p.add_run(part)
                p.style = 'Normal'
            else:
                # No inline code, handle as normal paragraph
                p = doc.add_paragraph(line, style='Normal')
        # Skip empty lines instead of adding empty paragraphs
        
        i += 1

if __name__ == "__main__":
    main()